"""Report generation for the standalone analyzer."""

from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Any

import numpy as np
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def generate_report(
    data: dict[str, Any],
    predictions: list[float] | np.ndarray,
    output_path: str | Path = "report.docx",
) -> None:
    """Generate a Word report with analysis results."""
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(12)

    title = doc.add_heading("聚乙烯补口粘接质量检测分析报告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("一、检测信息", level=1)
    doc.add_paragraph(f"报告生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"数据来源：{data.get('source', '未知')}")
    doc.add_paragraph(f"分析点数：{len(predictions)}")

    doc.add_heading("二、预测结果", level=1)
    preds = np.asarray(predictions)
    if len(preds) > 0:
        table = doc.add_table(rows=min(len(preds), 20) + 1, cols=2)
        table.style = "Light Grid Accent 1"
        table.rows[0].cells[0].text = "序号"
        table.rows[0].cells[1].text = "预测粘接力 (N/cm)"
        for i, p in enumerate(preds[:20]):
            table.rows[i + 1].cells[0].text = str(i + 1)
            table.rows[i + 1].cells[1].text = f"{float(p):.2f}"

        doc.add_heading("三、统计摘要", level=1)
        doc.add_paragraph(f"平均粘接力：{float(np.mean(preds)):.2f} N/cm")
        doc.add_paragraph(f"最大粘接力：{float(np.max(preds)):.2f} N/cm")
        doc.add_paragraph(f"最小粘接力：{float(np.min(preds)):.2f} N/cm")
        doc.add_paragraph(f"标准差：{float(np.std(preds)):.2f} N/cm")

    doc.add_heading("四、结论", level=1)
    doc.add_paragraph(
        "本报告由聚乙烯粘接声力耦合分析系统 V1.0 自动生成，"
        "分析结果仅供技术参考，最终判定需结合工程标准与专家意见。"
    )

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))

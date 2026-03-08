"""Word report generation using python-docx."""

import asyncio
from datetime import datetime
from pathlib import Path
from typing import Any, Optional

from core.config import get_settings

settings = get_settings()


class ReportGeneratorService:
    """Generate Word reports for defect analysis."""

    def __init__(self) -> None:
        self._output_dir = Path(settings.report_output_dir)

    async def analyze_image(self, image_bytes: bytes) -> dict[str, Any]:
        """Analyze C-scan image - placeholder for image-report-system integration."""
        # In production: call image-report-system defect analysis
        return {
            "defect_count": 0,
            "defects": [],
            "defect_area_ratio": 0.0,
            "mask_path": None,
        }

    async def generate_report(
        self,
        specimen_id: str,
        inspection_date: str,
        operator: str,
        equipment: str,
        defect_count: int,
        defect_area_ratio: float,
        predicted_force: float,
        overall_quality: str,
        image_path: Optional[str] = None,
    ) -> Path:
        """Generate Word report from analysis data."""
        return await asyncio.to_thread(
            self._generate_report_sync,
            specimen_id=specimen_id,
            inspection_date=inspection_date,
            operator=operator,
            equipment=equipment,
            defect_count=defect_count,
            defect_area_ratio=defect_area_ratio,
            predicted_force=predicted_force,
            overall_quality=overall_quality,
            image_path=image_path,
        )

    def _generate_report_sync(
        self,
        specimen_id: str,
        inspection_date: str,
        operator: str,
        equipment: str,
        defect_count: int,
        defect_area_ratio: float,
        predicted_force: float,
        overall_quality: str,
        image_path: Optional[str] = None,
    ) -> Path:
        from docx import Document
        from docx.shared import Inches, Pt
        doc = Document()
        doc.add_heading("聚乙烯补口粘接质量超声检测报告", 0)

        doc.add_heading("一、检测基本信息", level=1)
        table = doc.add_table(rows=6, cols=2)
        table.style = "Table Grid"
        rows_data = [
            ("试样编号", specimen_id),
            ("检测日期", inspection_date),
            ("操作人员", operator),
            ("检测设备", equipment),
            ("报告生成时间", datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
        ]
        for i, (k, v) in enumerate(rows_data):
            table.rows[i].cells[0].text = k
            table.rows[i].cells[1].text = str(v)

        doc.add_heading("二、检测结果总览", level=1)
        table2 = doc.add_table(rows=5, cols=2)
        table2.style = "Table Grid"
        results = [
            ("缺陷数量", str(defect_count)),
            ("缺陷面积占比", f"{defect_area_ratio:.2%}"),
            ("预测粘接力", f"{predicted_force:.1f} N/cm"),
            ("综合质量判定", overall_quality),
        ]
        for i, (k, v) in enumerate(results):
            table2.rows[i].cells[0].text = k
            table2.rows[i].cells[1].text = v

        doc.add_heading("三、结论与建议", level=1)
        if overall_quality == "合格":
            doc.add_paragraph(
                "本次检测结果表明，试样粘接质量满足标准要求，建议予以验收通过。"
            )
        elif overall_quality == "不合格":
            doc.add_paragraph(
                "本次检测结果表明，试样粘接质量未达到标准要求，建议进行返工处理。"
            )
        else:
            doc.add_paragraph("本次检测结果需人工复核确认，建议安排现场复检。")

        if image_path and Path(image_path).exists():
            doc.add_heading("四、缺陷分割结果", level=1)
            doc.add_picture(image_path, width=Inches(5))

        self._output_dir.mkdir(parents=True, exist_ok=True)
        filename = f"report_{specimen_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        out_path = self._output_dir / filename
        doc.save(str(out_path))
        return out_path

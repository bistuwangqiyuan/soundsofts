"""Word report template engine using python-docx."""

from __future__ import annotations

from pathlib import Path
from typing import TYPE_CHECKING

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

if TYPE_CHECKING:
    from ..pipeline import DetectedDefect


class ReportTemplate:
    """Generate structured Word reports with professional formatting."""

    def __init__(self) -> None:
        self.doc = Document()
        self._setup_styles()
        self._setup_page()

    def _setup_page(self) -> None:
        section = self.doc.sections[0]
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.0)

    def _setup_styles(self) -> None:
        style = self.doc.styles["Normal"]
        font = style.font
        font.name = "宋体"
        font.size = Pt(12)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.25

    def add_cover_page(self, title: str, subtitle: str, org: str, date: str, report_id: str) -> None:
        for _ in range(6):
            self.doc.add_paragraph("")

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        run.font.size = Pt(26)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

        p2 = self.doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(subtitle)
        run2.font.size = Pt(14)
        run2.font.color.rgb = RGBColor(0x63, 0x6E, 0x72)

        for _ in range(8):
            self.doc.add_paragraph("")

        info_lines = [
            f"报告编号：{report_id}",
            f"检测机构：{org}",
            f"报告日期：{date}",
        ]
        for line in info_lines:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line)
            run.font.size = Pt(13)

        self.add_page_break()

    def add_title(self, title: str) -> None:
        p = self.doc.add_heading(title, level=0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def add_section(self, heading: str, level: int = 1) -> None:
        self.doc.add_heading(heading, level=level)

    def add_paragraph(self, text: str, bold: bool = False, indent: bool = False) -> None:
        p = self.doc.add_paragraph()
        if indent:
            p.paragraph_format.first_line_indent = Cm(0.75)
        run = p.add_run(text)
        run.bold = bold

    def add_page_break(self) -> None:
        self.doc.add_page_break()

    def add_table(self, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
        table = self.doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = "Light Grid Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.bold = True

        for r_idx, row in enumerate(rows):
            for c_idx, cell_text in enumerate(row):
                table.rows[r_idx + 1].cells[c_idx].text = cell_text

    def add_defect_detail_table(self, defects: list[DetectedDefect]) -> None:
        headers = ["编号", "类型", "严重度", "面积(px)", "位置(x,y)", "置信度"]
        rows = []
        for d in defects:
            rows.append([
                str(d.defect_id),
                d.defect_type.value,
                d.severity.value,
                str(d.area_px),
                f"({d.centroid[0]:.0f}, {d.centroid[1]:.0f})",
                f"{d.confidence:.0%}",
            ])
        if not rows:
            rows = [["--"] * len(headers)]
        self.add_table(headers, rows)

    def add_image(self, image_path: str | Path, width_inches: float = 5.0, caption: str = "") -> None:
        p = Path(image_path)
        if p.exists():
            self.doc.add_picture(str(p), width=Inches(width_inches))
            if caption:
                cap = self.doc.add_paragraph()
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = cap.add_run(caption)
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x63, 0x6E, 0x72)

    def add_header_footer(self, report_id: str, org: str) -> None:
        section = self.doc.sections[0]
        header = section.header
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.text = f"{org}    报告编号: {report_id}"
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in hp.runs:
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        footer = section.footer
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.text = "本报告由超声图像检测与报告生成系统自动生成"
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in fp.runs:
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    def save(self, output_path: str | Path) -> None:
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(output_path))
